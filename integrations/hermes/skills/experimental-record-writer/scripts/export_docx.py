#!/usr/bin/env python3
"""Export the current experimental-record bundle as a table-oriented DOCX."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

from record_bundle import (
    DEFAULT_PREVIOUS_RECORD,
    RecordError,
    append_audit,
    as_list,
    load_bundle,
    normalized_date,
    now_iso,
    sha256,
    text as record_text,
    write_manifest,
)

LATIN_FONT = "Arial"
CJK_FONT = "Arial Unicode MS"


def text(value: Any, missing: str = "未提供") -> str:
    return record_text(value, missing)


def display_date(value: Any) -> str:
    try:
        return normalized_date(value)
    except RecordError:
        return text(value)

def load_docx() -> Any:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION_START
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError as error:
        raise RecordError("缺少 python-docx。请使用带 python-docx 的 Python 运行时后重试。") from error
    return Document, WD_SECTION_START, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt


def set_cell(cell: Any, value: Any, bold: bool = False) -> None:
    from docx.shared import Pt

    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text(value))
    run.bold = bold
    run.font.size = Pt(9)


def add_table(
    document: Any, headers: list[str], rows: list[list[Any]], column_widths: list[float] | None = None
) -> None:
    from docx.oxml import OxmlElement

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = column_widths is None
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell(cell, header, bold=True)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "true")
    header_properties.append(repeat_header)
    for row in rows:
        table_row = table.add_row()
        row_properties = table_row._tr.get_or_add_trPr()
        keep_row_together = OxmlElement("w:cantSplit")
        row_properties.append(keep_row_together)
        cells = table_row.cells
        for cell, value in zip(cells, row):
            set_cell(cell, value)
    if column_widths:
        if len(column_widths) != len(headers):
            raise RecordError("DOCX 表格列宽数量必须与表头数量一致")
        from docx.shared import Inches

        for column, width in zip(table.columns, column_widths):
            column.width = Inches(width)
        for table_row in table.rows:
            for cell, width in zip(table_row.cells, column_widths):
                cell.width = Inches(width)
    document.add_paragraph()


def values(items: Any, key: str = "name") -> list[dict[str, Any]]:
    return [item if isinstance(item, dict) else {key: item} for item in as_list(items)]


def is_image(path: Path) -> bool:
    return path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}


def set_east_asia_font(document: Any, oxml_element: Any, qn: Any) -> None:
    """Set a macOS/LibreOffice-compatible CJK fallback on every concrete run."""
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            properties = run._element.get_or_add_rPr()
            fonts = properties.rFonts
            if fonts is None:
                fonts = oxml_element("w:rFonts")
                properties.insert(0, fonts)
            run.font.name = LATIN_FONT
            fonts.set(qn("w:ascii"), LATIN_FONT)
            fonts.set(qn("w:hAnsi"), LATIN_FONT)
            fonts.set(qn("w:eastAsia"), CJK_FONT)
            fonts.set(qn("w:cs"), LATIN_FONT)
            fonts.set(qn("w:hint"), "eastAsia")
            language = properties.find(qn("w:lang"))
            if language is None:
                language = oxml_element("w:lang")
                properties.append(language)
            language.set(qn("w:val"), "zh-CN")
            language.set(qn("w:eastAsia"), "zh-CN")


def add_heading(document: Any, title: str, level: int = 1) -> None:
    # Avoid Word's themed Heading styles: LibreOffice can select a non-CJK
    # theme font for them even when a run has an East Asian fallback.
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"


def export_document(directory: Path, output: Path) -> None:
    Document, _, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt = load_docx()
    bundle = load_bundle(directory)
    record = bundle["record"]
    if output.exists():
        raise RecordError(f"目标 DOCX 已存在，拒绝覆盖: {output}")

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("实验记录")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(text(record.get("title"))).bold = True
    document.add_paragraph()

    add_table(document, ["字段", "内容"], [
        ["记录 ID", record.get("id")],
        ["修订", record.get("revision")],
        ["实际执行日期", record.get("performedAt")],
        ["执行者", ", ".join(text(item, "") for item in as_list(record.get("performedBy"))) or "待确认"],
        ["项目", record.get("project")],
        ["创建日期", display_date(record.get("createdAt"))],
    ])

    add_heading(document, "目的")
    document.add_paragraph(text(record.get("objective")))

    technique = record.get("technique", {})
    protocol = record.get("protocol", {})
    technique_rows = [["技术名称", technique.get("name")], ["方案", protocol.get("title")], ["方案版本", protocol.get("version")]]
    if any(text(row[1]) != "未提供" for row in technique_rows):
        add_heading(document, "技术与方案")
        add_table(document, ["字段", "内容"], technique_rows)

    add_heading(document, "样本与输入")
    samples = values(record.get("samples"), "id")
    if samples:
        add_table(document, ["样本", "描述"], [[item.get("id"), item.get("description")] for item in samples])
    else:
        document.add_paragraph("待确认样本信息。")

    groups = [item for item in as_list(record.get("groups")) if isinstance(item, dict)]
    if groups:
        add_heading(document, "实验分组")
        add_table(document, ["分组", "样本/细胞系", "检测目标"], [
            [item.get("id"), item.get("sample"), "；".join(text(target, "") for target in as_list(item.get("targets")))]
            for item in groups
        ])

    add_heading(document, "试剂", level=2)
    reagents = values(record.get("reagents"))
    if reagents:
        add_table(document, ["名称", "厂家", "货号", "类别", "浓度/稀释度"], [[
            item.get("reagentName") or item.get("name"),
            item.get("manufacturer") or item.get("vendor"),
            item.get("catalogNumber") or item.get("catalogNo"),
            item.get("category"),
            item.get("concentration"),
        ] for item in reagents], [1.25, 1.25, 1.15, 1.1, 1.25])
    else:
        document.add_paragraph("待补充试剂信息。")

    add_heading(document, "仪器与软件", level=2)
    instruments = values(record.get("instruments"))
    software = values(record.get("software"))
    add_table(document, ["类别", "名称", "备注"], [
        *[["仪器", item.get("name"), item.get("configuration") or DEFAULT_PREVIOUS_RECORD] for item in instruments],
        *[["软件", item.get("name"), item.get("purpose") or DEFAULT_PREVIOUS_RECORD] for item in software],
    ])

    add_heading(document, "计划步骤")
    planned_steps = as_list(record.get("plannedSteps"))
    if planned_steps:
        for item in planned_steps:
            document.add_paragraph(text(item), style="List Number")
    else:
        document.add_paragraph("未提供计划步骤。")

    add_heading(document, "实际执行")
    actual_steps = values(record.get("actualSteps"), "action")
    if actual_steps:
        rows = []
        for item in actual_steps:
            parameters = []
            for parameter in as_list(item.get("parameters")):
                if isinstance(parameter, dict):
                    parameters.append(f"{text(parameter.get('name'))}={text(parameter.get('value'))}{text(parameter.get('unit'), '')}")
                else:
                    parameters.append(text(parameter))
            rows.append([
                item.get("sequence"),
                item.get("action"),
                "; ".join(parameters) or DEFAULT_PREVIOUS_RECORD,
                item.get("performedBy") or ", ".join(text(person, "") for person in as_list(record.get("performedBy"))) or "待确认",
            ])
        add_table(document, ["序号", "实际操作", "关键参数", "执行者"], rows)
    else:
        document.add_paragraph("未提供实际执行步骤。")

    add_heading(document, "偏差与异常")
    deviations = values(record.get("deviations"), "description")
    if deviations:
        add_table(document, ["描述", "原因", "影响", "处置"], [[item.get("description"), item.get("cause"), item.get("impact"), item.get("action")] for item in deviations])

    add_heading(document, "结果与附件")
    results = [item for item in as_list(record.get("results")) if isinstance(item, dict)]
    if results:
        add_table(document, ["日期", "结果摘要", "记录人"], [[display_date(item.get("observedAt")), item.get("summary"), item.get("addedBy")] for item in results])
    else:
        document.add_paragraph("未提供结果。")

    attachments = {item.get("id"): item for item in bundle.get("attachments", []) if isinstance(item, dict)}
    if attachments:
        add_heading(document, "附件清单")
        for attachment in attachments.values():
            source = directory / "attachments" / str(attachment.get("storedName", ""))
            document.add_paragraph(f"{attachment.get('id')}: {attachment.get('originalName')} | SHA-256: {attachment.get('sha256')}")
            if source.is_file() and is_image(source):
                try:
                    document.add_picture(str(source), width=Inches(5.8))
                    caption = document.add_paragraph(f"图片：{attachment.get('originalName')}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as error:  # python-docx supports only a subset of image encodings
                    document.add_paragraph(f"未嵌入图片，请查看原始附件：{error}")

    if text(record.get("conclusion")) != "未提供":
        add_heading(document, "结论")
        document.add_paragraph(text(record.get("conclusion")))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{record.get('id')} | 修订 {record.get('revision')} | 导出日期 {display_date(now_iso())}")
    set_east_asia_font(document, OxmlElement, qn)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    export = {"path": str(output.relative_to(directory)), "sha256": sha256(output), "revision": record["revision"], "format": "docx"}
    write_manifest(directory, bundle, [export])
    append_audit(directory, "EXPORT_DOCX", bundle, "未提供", "导出 DOCX 快照", {"path": export["path"], "sha256": export["sha256"]})


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--record", required=True, help="record-bundle directory")
    parser.add_argument("--out", help="optional DOCX output path inside the record directory")
    args = parser.parse_args()
    try:
        directory = Path(args.record).expanduser().resolve()
        bundle = load_bundle(directory)
        revision = int(bundle["record"]["revision"])
        output = Path(args.out).expanduser().resolve() if args.out else directory / "exports" / f"record-r{revision:04d}.docx"
        try:
            output.relative_to(directory)
        except ValueError as error:
            raise RecordError("DOCX 输出必须位于记录目录内") from error
        export_document(directory, output)
        print(output)
    except RecordError as error:
        parser.error(str(error))


if __name__ == "__main__":
    main()
